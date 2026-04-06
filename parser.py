import fitz
import docx
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import os
import zipfile


# ─────────────────────────────────────────────
# SUPPORTED FORMATS
# ─────────────────────────────────────────────
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".png", ".jpg", ".jpeg"}


# ─────────────────────────────────────────────
# FILE VALIDATION
# ─────────────────────────────────────────────
def validate_file(file_path):
    """
    Returns a dict:
      { "valid": True/False, "reason": "..." }
    Reasons: "unsupported_format" | "corrupt" | "empty" | "ok"
    """
    ext = os.path.splitext(file_path)[1].lower()

    # 1. Check extension
    if ext not in SUPPORTED_EXTENSIONS:
        return {"valid": False, "reason": "unsupported_format"}

    # 2. Check file is not empty
    if os.path.getsize(file_path) == 0:
        return {"valid": False, "reason": "empty"}

    # 3. Format-specific corruption checks
    try:
        if ext == ".pdf":
            pdf = fitz.open(file_path)
            if pdf.page_count == 0:
                return {"valid": False, "reason": "corrupt"}
            pdf.close()

        elif ext == ".docx":
            # DOCX is a ZIP — try opening it
            if not zipfile.is_zipfile(file_path):
                return {"valid": False, "reason": "corrupt"}
            docx.Document(file_path)  # full open check

        elif ext in {".png", ".jpg", ".jpeg"}:
            img = Image.open(file_path)
            img.verify()  # detects truncated/corrupt images

    except Exception:
        return {"valid": False, "reason": "corrupt"}

    return {"valid": True, "reason": "ok"}


# ─────────────────────────────────────────────
# PDF — ADAPTIVE LAYOUT-AWARE EXTRACTION
# ─────────────────────────────────────────────
def _detect_column_threshold(blocks, page_width):
    """
    Dynamically detect the x-midpoint that separates columns.
    Falls back to page centre if no clear gap is found.
    """
    x_starts = sorted(set(round(b[0] / 10) * 10 for b in blocks))

    # Find the largest gap in x-start positions
    best_gap = 0
    threshold = page_width / 2

    for i in range(1, len(x_starts)):
        gap = x_starts[i] - x_starts[i - 1]
        if gap > best_gap and x_starts[i - 1] > page_width * 0.2:
            best_gap = gap
            threshold = (x_starts[i] + x_starts[i - 1]) / 2

    return threshold


def _extract_page_text(page):
    """
    Smart per-page extraction:
    - Detects single-column vs multi-column layouts dynamically
    - Handles sidebar blocks (narrow left strip)
    - Falls back to raw text on error
    """
    try:
        page_width = page.rect.width
        blocks = page.get_text("blocks")  # (x0, y0, x1, y1, text, block_no, block_type)

        if not blocks:
            return page.get_text()

        # Filter out image blocks (type 1) — keep text blocks (type 0)
        text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]

        if not text_blocks:
            return page.get_text()

        # Detect if layout is multi-column
        x_positions = [b[0] for b in text_blocks]
        x_min, x_max = min(x_positions), max(x_positions)
        spread = x_max - x_min

        # If blocks are spread more than 35% of page width → multi-column
        if spread > page_width * 0.35:
            threshold = _detect_column_threshold(text_blocks, page_width)

            # Sidebar: very narrow left column (< 20% of page width)
            sidebar_limit = page_width * 0.20

            sidebar_blocks = [b for b in text_blocks if b[0] < sidebar_limit]
            left_blocks = [b for b in text_blocks if sidebar_limit <= b[0] < threshold]
            right_blocks = [b for b in text_blocks if b[0] >= threshold]

            sidebar_blocks = sorted(sidebar_blocks, key=lambda b: (b[1], b[0]))
            left_blocks = sorted(left_blocks, key=lambda b: (b[1], b[0]))
            right_blocks = sorted(right_blocks, key=lambda b: (b[1], b[0]))

            ordered = sidebar_blocks + left_blocks + right_blocks
        else:
            # Single column — sort top to bottom
            ordered = sorted(text_blocks, key=lambda b: (b[1], b[0]))

        page_text = ""
        for b in ordered:
            block_text = b[4].strip()
            if block_text:
                page_text += block_text + "\n"

        return page_text

    except Exception:
        return page.get_text()


def extract_pdf(file_path):
    text = ""
    try:
        pdf = fitz.open(file_path)
        for page in pdf:
            text += _extract_page_text(page) + "\n"
        pdf.close()
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {e}")
    return text


# ─────────────────────────────────────────────
# DOCX — FULL EXTRACTION (paragraphs + tables)
# ─────────────────────────────────────────────
def extract_docx(file_path):
    try:
        doc = docx.Document(file_path)
        parts = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1]  # strip namespace

            if tag == "p":
                # Paragraph
                para_text = "".join(
                    run.text for run in element.iterchildren()
                    if run.tag.split("}")[-1] == "r"
                    and hasattr(run, "text")
                ).strip()

                # Use python-docx paragraph for cleaner text
                pass

            elif tag == "tbl":
                # Table — extract cell text row by row
                for row in element.iterchildren():
                    if row.tag.split("}")[-1] != "tr":
                        continue
                    row_cells = []
                    for cell in row.iterchildren():
                        if cell.tag.split("}")[-1] == "tc":
                            cell_text = " ".join(
                                p.text_content() if hasattr(p, "text_content")
                                else "".join(r.text or "" for r in p.iterchildren())
                                for p in cell.iterchildren()
                            ).strip()
                            if cell_text:
                                row_cells.append(cell_text)
                    if row_cells:
                        parts.append(" | ".join(row_cells))

        # Cleaner paragraph extraction via python-docx API
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # Merge tables and paragraphs
        # Use paragraph text as base, append table parts
        combined = "\n".join(full_text)
        if parts:
            combined += "\n" + "\n".join(parts)

        return combined

    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {e}")


# ─────────────────────────────────────────────
# IMAGE — ENHANCED OCR
# ─────────────────────────────────────────────
def extract_image(file_path):
    try:
        img = Image.open(file_path).convert("L")  # greyscale

        # Enhance contrast for better OCR
        img = ImageEnhance.Contrast(img).enhance(2.0)
        img = img.filter(ImageFilter.SHARPEN)

        # Upscale small images
        w, h = img.size
        if w < 1000:
            scale = 1000 / w
            img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

        # OCR with layout-aware config
        custom_config = r"--oem 3 --psm 6"
        text = pytesseract.image_to_string(img, config=custom_config)
        return text

    except Exception as e:
        raise RuntimeError(f"Image OCR failed: {e}")


# ─────────────────────────────────────────────
# MAIN ROUTER
# ─────────────────────────────────────────────
def extract_text(file_path):
    """
    Validates and extracts text from a file.
    Returns: (text: str, status: dict)
      status = { "valid": bool, "reason": str }
    """
    status = validate_file(file_path)

    if not status["valid"]:
        return "", status

    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            text = extract_pdf(file_path)
        elif ext == ".docx":
            text = extract_docx(file_path)
        elif ext in {".png", ".jpg", ".jpeg"}:
            text = extract_image(file_path)
        else:
            return "", {"valid": False, "reason": "unsupported_format"}

        if not text.strip():
            return "", {"valid": False, "reason": "empty"}

        return text, {"valid": True, "reason": "ok"}

    except RuntimeError as e:
        return "", {"valid": False, "reason": f"corrupt — {e}"}
